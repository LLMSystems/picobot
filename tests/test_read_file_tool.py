from pathlib import Path

from simplified_chatbot.tools.filesystem import ReadFileTool, build_default_tool_registry
from simplified_chatbot.tools.registry import ToolRegistry


def test_read_file_basic_read_has_line_numbers(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("alpha\nbeta\ngamma\n", encoding="utf-8")
    registry = ToolRegistry()
    registry.register(ReadFileTool(workspace=tmp_path, allowed_dir=tmp_path))

    result = registry.execute("read_file", {"path": str(file_path)})

    assert "1| alpha" in result
    assert "3| gamma" in result
    assert "End of file" in result


def test_read_file_offset_and_limit(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("\n".join(f"line {i}" for i in range(1, 8)), encoding="utf-8")
    registry = ToolRegistry()
    registry.register(ReadFileTool(workspace=tmp_path, allowed_dir=tmp_path))

    result = registry.execute(
        "read_file",
        {"path": str(file_path), "offset": 3, "limit": 2},
    )

    assert "3| line 3" in result
    assert "4| line 4" in result
    assert "5| line 5" not in result
    assert "Use offset=5 to continue" in result


def test_read_file_returns_dedup_stub_on_second_read(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("hello\nworld\n", encoding="utf-8")
    registry = build_default_tool_registry(workspace=tmp_path)

    first = registry.execute("read_file", {"path": str(file_path)})
    second = registry.execute("read_file", {"path": str(file_path)})

    assert "1| hello" in first
    assert "[File unchanged since last read:" in second


def test_read_file_blocks_paths_outside_workspace(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    outside = tmp_path / "outside"
    outside.mkdir()
    file_path = outside / "secret.txt"
    file_path.write_text("secret", encoding="utf-8")

    registry = build_default_tool_registry(workspace=workspace)
    result = registry.execute("read_file", {"path": str(file_path)})

    assert "Error:" in result
    assert "outside allowed directory" in result


def _write_simple_pdf(path: Path, text: str) -> None:
    content = "\n".join(["BT", "/F1 18 Tf", "72 120 Td", f"({text}) Tj", "ET"]).encode("utf-8")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode("ascii"))
        result.extend(obj)
        result.extend(b"\nendobj\n")
    xref_offset = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    result.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii"),
    )
    path.write_bytes(result)


def _read(tmp_path: Path, args: dict) -> str:
    registry = ToolRegistry()
    registry.register(ReadFileTool(workspace=tmp_path, allowed_dir=tmp_path))
    return registry.execute("read_file", args)


def test_read_file_dispatches_pdf_with_pages(tmp_path: Path):
    _write_simple_pdf(tmp_path / "sample.pdf", "Hello PDF")

    result = _read(tmp_path, {"path": "sample.pdf", "pages": "1"})

    assert "Hello PDF" in result
    assert "[Page 1]" in result


def test_read_file_dispatches_docx(tmp_path: Path):
    from docx import Document

    document = Document()
    document.add_paragraph("Routed through read_file.")
    document.save(tmp_path / "notes.docx")

    result = _read(tmp_path, {"path": "notes.docx"})

    assert "Routed through read_file." in result
    assert "[Paragraph 1]" in result


def test_read_file_dispatches_xlsx_with_sheet_and_range(tmp_path: Path):
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Main"
    sheet["A1"] = "key"
    sheet["B1"] = "value"
    sheet["A2"] = "count"
    sheet["B2"] = 7
    workbook.save(tmp_path / "data.xlsx")

    result = _read(tmp_path, {"path": "data.xlsx", "sheet": "Main", "range": "A1:B2"})

    assert "[Sheet: Main]" in result
    assert "1| key\tvalue" in result
    assert "2| count\t7" in result


def test_read_file_office_dispatch_ignores_line_pagination(tmp_path: Path):
    from docx import Document

    document = Document()
    document.add_paragraph("Still routed via extraction.")
    document.save(tmp_path / "doc.docx")

    # offset/limit are text-only; office dispatch must ignore them, not error.
    result = _read(tmp_path, {"path": "doc.docx", "offset": 5, "limit": 1})

    assert "Still routed via extraction." in result
    assert "Error" not in result


def test_read_file_rejects_non_office_binary(tmp_path: Path):
    (tmp_path / "blob.bin").write_bytes(b"\x00\x01\x02\xff\xfe")

    result = _read(tmp_path, {"path": "blob.bin"})

    assert "Cannot read binary file" in result
