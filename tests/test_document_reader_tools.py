import asyncio
from pathlib import Path

from simplified_chatbot.tools.document_readers import (
    ReadDocxTool,
    ReadPdfTool,
    ReadXlsxTool,
)


def _write_simple_pdf(path: Path, text: str) -> None:
    content = "\n".join(
        [
            "BT",
            "/F1 18 Tf",
            "72 120 Td",
            f"({text}) Tj",
            "ET",
        ],
    ).encode("utf-8")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode("ascii"))
        result.extend(obj)
        result.extend(b"\nendobj\n")

    xref_offset = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    result.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii"),
    )
    path.write_bytes(result)


def test_read_pdf_tool_extracts_text(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    _write_simple_pdf(pdf_path, "Hello PDF")
    tool = ReadPdfTool(workspace=tmp_path, allowed_dir=tmp_path)

    result = asyncio.run(tool.execute(path="sample.pdf"))

    assert "Hello PDF" in result
    assert "[Page 1]" in result


def test_read_docx_tool_extracts_paragraphs_and_tables(tmp_path):
    from docx import Document

    docx_path = tmp_path / "plan.docx"
    document = Document()
    document.add_heading("Plan", level=1)
    document.add_paragraph("Ship the first version.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Ready"
    document.save(docx_path)

    tool = ReadDocxTool(workspace=tmp_path, allowed_dir=tmp_path)
    result = asyncio.run(tool.execute(path="plan.docx"))

    assert "Plan" in result
    assert "Ship the first version." in result
    assert "[Table 1]" in result
    assert "Status | Ready" in result


def test_read_xlsx_tool_extracts_rows(tmp_path):
    from openpyxl import Workbook

    xlsx_path = tmp_path / "report.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet["A1"] = "Name"
    sheet["B1"] = "Value"
    sheet["A2"] = "Revenue"
    sheet["B2"] = 42
    workbook.save(xlsx_path)

    tool = ReadXlsxTool(workspace=tmp_path, allowed_dir=tmp_path)
    result = asyncio.run(tool.execute(path="report.xlsx", sheet="Summary", range="A1:B2"))

    assert "[Sheet: Summary]" in result
    assert "[Range: A1:B2]" in result
    assert "1| Name\tValue" in result
    assert "2| Revenue\t42" in result
